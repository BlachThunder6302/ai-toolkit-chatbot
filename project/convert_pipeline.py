import os
import json
import re
import pdfplumber
from docx import Document
from docx.shared import Pt
import pypandoc
import fitz  # pymupdf for images

# -----------------------
# CONFIG
# -----------------------

INPUT_DIR = "input_pdfs"
DOCX_DIR = "out_docx"
MD_DIR = "out_md"
MEDIA_DIR = "media"
PROV_DIR = "provenance"

os.makedirs(DOCX_DIR, exist_ok=True)
os.makedirs(MD_DIR, exist_ok=True)
os.makedirs(MEDIA_DIR, exist_ok=True)
os.makedirs(PROV_DIR, exist_ok=True)

# -----------------------
# HELPERS
# -----------------------

def sanitize_name(name: str) -> str:
    return re.sub(r"[^A-Za-z0-9._-]+", "_", name)


def clean_text(text: str) -> str:
    """
    Option A: Minimal cleanup only.
    Does NOT restructure or modify meaning.
    Only removes noise introduced by extraction.
    """
    if not text:
        return ""

    # Normalize whitespace
    text = re.sub(r"\s+", " ", text).strip()

    # Remove numeric garbage from PDF headers/footers/OCR noise
    # e.g. lines like: 3995991, 548995, 1318513
    text = re.sub(r"\b\d{5,}\b", "", text)

    # Remove isolated zero lines
    text = re.sub(r"\b0\b", "", text)

    # Remove duplicated words (light touch)
    text = re.sub(r"\b(\w+)\s+\1\b", r"\1", text)

    return text.strip()


# -----------------------
# TEXT EXTRACTION
# -----------------------

def extract_pages_as_text(pdf_path):
    pages = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            txt = page.extract_text() or ""
            # Fix hyphen line-breaks then collapse newlines
            txt = txt.replace("-\n", "")
            txt = txt.replace("\n", " ")

            txt = clean_text(txt)
            pages.append(txt)
    return pages


# -----------------------
# IMAGE EXTRACTION
# -----------------------

def extract_images(pdf_path, doc_slug):
    img_index = {}
    target_dir = os.path.join(MEDIA_DIR, doc_slug)
    os.makedirs(target_dir, exist_ok=True)

    doc = fitz.open(pdf_path)

    for pno in range(len(doc)):
        paths = []
        page = doc[pno]

        for imgid, img in enumerate(page.get_images(full=True), start=1):
            xref = img[0]
            pix = fitz.Pixmap(doc, xref)

            if pix.alpha:
                try:
                    pix = fitz.Pixmap(fitz.csRGB, pix)
                except:
                    continue

            out_path = os.path.join(
                target_dir,
                f"page-{pno+1}-img-{imgid}.png"
            )

            try:
                pix.save(out_path)
                paths.append(os.path.relpath(out_path))
            except Exception:
                print(f"Skipping image on page {pno+1}, xref {xref}: unsupported colorspace")

        if paths:
            img_index[pno+1] = paths

    doc.close()
    return img_index


# -----------------------
# BUILD DOCX WITH PAGE MARKERS
# -----------------------

def build_docx_with_page_headings(pdf_path, docx_path, insert_image_placeholders=True):

    pages = extract_pages_as_text(pdf_path)
    doc = Document()

    # Title
    title = os.path.splitext(os.path.basename(pdf_path))[0]
    doc.add_heading(title, level=1)

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Extract images
    doc_slug = sanitize_name(title)
    img_index = extract_images(pdf_path, doc_slug)

    prov = {"title": title, "pages": []}

    # Write page headings + text
    for i, text in enumerate(pages, start=1):
        doc.add_heading(f"Page {i}", level=2)

        if text.strip():
            doc.add_paragraph(text)
        else:
            doc.add_paragraph("[No extractable text on this page]")

        # Optional image placeholders
        if insert_image_placeholders and img_index.get(i):
            doc.add_paragraph("Images on this page:")
            for rel in img_index[i]:
                doc.add_paragraph(f"[image]: {rel}")

        prov["pages"].append({
            "page": i,
            "char_count": len(text),
            "has_images": bool(img_index.get(i, []))
        })

    # Save DOCX
    doc.save(docx_path)

    # Save provenance JSON
    prov_path = os.path.join(PROV_DIR, f"{doc_slug}.pages.json")
    with open(prov_path, "w", encoding="utf-8") as f:
        json.dump(prov, f, ensure_ascii=False, indent=2)

    return len(pages), prov_path


# -----------------------
# DOCX → Markdown (Pandoc)
# -----------------------

def docx_to_markdown(docx_path, md_path, media_out_dir):
    os.makedirs(media_out_dir, exist_ok=True)

    pypandoc.convert_file(
        docx_path,
        "gfm",
        outputfile=md_path,
        extra_args=[
            "--wrap=none",
            f"--extract-media={media_out_dir}",
        ],
    )

# -----------------------
# MAIN PIPELINE
# -----------------------

def process_all_pdfs():
    for fname in os.listdir(INPUT_DIR):
        if not fname.lower().endswith(".pdf"):
            continue

        pdf_path = os.path.join(INPUT_DIR, fname)
        base = sanitize_name(os.path.splitext(fname)[0])

        docx_path = os.path.join(DOCX_DIR, f"{base}.docx")
        md_path   = os.path.join(MD_DIR,   f"{base}.md")
        media_out = os.path.join(MEDIA_DIR, base)

        print(f"\n→ Converting: {fname}")

        pages, prov = build_docx_with_page_headings(pdf_path, docx_path)
        print(f"   DOCX saved: {docx_path} ({pages} pages)")
        print(f"   Provenance saved: {prov}")

        docx_to_markdown(docx_path, md_path, media_out)
        print(f"   Markdown saved: {md_path}")
        print(f"   Media saved to: {media_out}")


if __name__ == "__main__":
    process_all_pdfs()
    print("\n✅ DONE (Option A Cleanup Applied)")
